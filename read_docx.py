import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading DOCX: {e}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <path_to_docx>")
    else:
        print(read_docx(sys.argv[1]))
